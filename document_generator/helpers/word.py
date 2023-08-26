from docx.shared import Pt

from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches
from docx.enum.text import WD_TAB_ALIGNMENT
import docx
from typing import Optional


class WordHelper:
    def __init__(self, doc: docx.Document) -> None:
        self.document = doc

    def add_user_input_with_topics(
        self, topic_name: str, user_input: str, topic_bold: bool, topic_font_size: int
    ) -> docx.text.paragraph.Paragraph:
        paragraph = self.document.add_paragraph()
        topic_name_run = paragraph.add_run(topic_name)
        topic_name_run.font.size = Pt(topic_font_size)
        if topic_bold:
            topic_name_run.font.bold = True

        run_for_colon = paragraph.add_run(": ")
        run_for_colon.font.size = Pt(13)
        run_for_colon.font.bold = True

        user_input_run = paragraph.add_run(user_input)
        user_input_run.font.size = Pt(12)

        return paragraph

    def add_paragraph(
        self,
        text="",
        alignment=None,
        font_size=None,
        font_bold=False,
    ) -> docx.text.paragraph.Paragraph:
        paragraph = self.document.add_paragraph()

        if alignment == "center":
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        if not text:
            return paragraph

        run = paragraph.add_run(text)
        run.font.name = "Times New Roman"

        if font_size:
            run.font.size = Pt(font_size)
        if font_bold:
            run.font.bold = True

        return paragraph

    def set_ruler_position(self, paragraph, position) -> None:
        """
        Use like :
            self.set_ruler_position(paragraph, ruler_position)
        """
        tab_stops = paragraph.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Inches(position), WD_TAB_ALIGNMENT.LEFT)

    def set_paragraph_line_spacing(self, paragraph, line_spacing) -> None:
        """
        Set custom line spacing for a paragraph.
        :param paragraph: The paragraph for which line spacing is to be set.
        :param line_spacing: Line spacing value in points.
        """
        paragraph_format = paragraph.paragraph_format
        paragraph_format.line_spacing = line_spacing
        # Optional: Adjust the spacing before the paragraph.
        paragraph_format.space_before = Pt(0)
        # Optional: Adjust the spacing after the paragraph.
        paragraph_format.space_after = Pt(0)
