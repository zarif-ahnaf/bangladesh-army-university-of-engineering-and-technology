import docx
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches
import io
import base64

from ..helpers.word import WordHelper

from ..resources.logo import BAUET_LOGO


class BuildUniversityLogoAndName:
    def __init__(self, doc: docx.Document) -> None:
        self.document = doc
        self.word_helper = WordHelper(doc)

    def __build_university_name__(self) -> None:
        self.word_helper.add_paragraph(
            text="Bangladesh Army University of Engineering & Technology (BAUET)",
            alignment="center",
            font_size=14,
            font_bold=True,
        )

    def __build_university_location__(self) -> None:
        self.word_helper.add_paragraph(
            text="Qadirabad Cantonment-6431, Natore",
            alignment="center",
            font_size=11,
        )

    def __build_university_logo__(
        self, image_path: str, width: float, height: float
    ) -> None:
        paragraph = self.word_helper.add_paragraph()
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        run = paragraph.add_run()
        run.add_picture(image_path, width=Inches(width), height=Inches(height))

    def build(self) -> docx.Document:
        # Builder methods
        self.__build_university_name__()
        self.__build_university_location__()
        self.__build_university_logo__(
            io.BytesIO(base64.b64decode(BAUET_LOGO)), 1.5, 1.5
        )
        # Empty
        self.word_helper.add_paragraph()
