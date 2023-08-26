from ._factory import BaseFactory
from .resources.data import DATA
from docx.shared import Pt

from .builders.build_university_logo_and_name import BuildUniversityLogoAndName


class Builder(BaseFactory):
    def __build_course_name__(self, course_name: str) -> None:
        self.word_helper.add_user_input_with_topics(
            topic_name="Course No.",
            user_input=course_name,
            topic_bold=True,
            topic_font_size=14,
        )

    def __build_course_title__(self, course_title: str) -> None:
        self.word_helper.add_user_input_with_topics(
            topic_name="Course Name",
            user_input=course_title,
            topic_bold=True,
            topic_font_size=14,
        )

    def __build_date_of_experiment__(self, date_of_experiment: str) -> None:
        self.word_helper.add_user_input_with_topics(
            topic_name="Date of Experiment",
            user_input=date_of_experiment,
            topic_bold=True,
            topic_font_size=14,
        )

    def __build_date_of_submission__(self, date_of_submission: str) -> None:
        self.word_helper.add_user_input_with_topics(
            topic_name="Date of Submission",
            user_input=date_of_submission,
            topic_bold=True,
            topic_font_size=14,
        )

    def __build_experiment_no__(self, experiment_no: str) -> None:
        self.word_helper.add_user_input_with_topics(
            topic_name="Experiment No.",
            user_input=experiment_no,
            topic_bold=True,
            topic_font_size=14,
        )

    def __build_experiment_name__(self, experiment_name: str) -> None:
        self.word_helper.add_user_input_with_topics(
            topic_name="Experiment Name",
            user_input=experiment_name,
            topic_bold=True,
            topic_font_size=14,
        )

    def build(
        self,
        course_name: str,
        course_title: str,
        date_of_experiment: str,
        date_of_submission: str,
        experiment_no: str,
        experiment_name: str,
        # Start Student
        student_information: dict,
        key: str,
    ) -> None:
        data = DATA[key]
        BuildUniversityLogoAndName(self.document).build()

        self.__build_course_name__(course_name)
        self.__build_course_title__(course_title)

        self.word_helper.add_paragraph()

        self.__build_date_of_experiment__(date_of_experiment)
        self.__build_date_of_submission__(date_of_submission)

        self.word_helper.add_paragraph()

        self.__build_experiment_no__(experiment_no)
        self.__build_experiment_name__(experiment_name)

        self.word_helper.add_paragraph()

        teacher_data_list = [
            (key, value) for dictionary in data for key, value in dictionary.items()
        ]

        # Submitted By
        paragraph = self.word_helper.add_paragraph()
        left_side_item_list = []

        self.word_helper.set_paragraph_line_spacing(paragraph, Pt(20))
        submitted_by_run = paragraph.add_run("Submitted By")
        submitted_by_run.font.size = Pt(13)
        submitted_by_run.font.bold = True

        for i in range(
            0, max(len(teacher_data_list), len(student_information.items()))
        ):
            try:
                topic_name, user_input = list(student_information.items())[i]
                _paragraph = self.word_helper.add_user_input_with_topics(
                    str(topic_name), str(user_input), False, 11
                )

            except KeyError:
                _paragraph = self.word_helper.add_paragraph()

            finally:
                self.word_helper.set_paragraph_line_spacing(_paragraph, Pt(20))
                left_side_item_list.append(_paragraph)

        # Submitted To
        self.word_helper.set_ruler_position(paragraph, 4.0)
        paragraph.add_run("\t")
        submitted_to_run = paragraph.add_run("Submitted To")
        submitted_to_run.font.size = Pt(13)
        submitted_to_run.font.bold = True

        for index, (item, value) in enumerate(teacher_data_list):
            student_item = left_side_item_list[index]

            # Set ruler Postion
            self.word_helper.set_ruler_position(student_item, 4.0)
            student_item.add_run("\t")

            teacher_information = student_item.add_run(value)
            teacher_information.font.size = Pt(12)

            if item.lower() == "name":
                teacher_information.font.bold = True

        return self.document
